#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Application Flask simple pour piloter une détection de fraudes
en utilisant exclusivement le script local `fraude_detect_admi.py`.

- Page d'accueil: procédure et bouton Continuer
- Page /dates: saisie des dates (validation serveur)
- Lancement du traitement: exécution du script en sous-processus avec la fenêtre
  temporelle transmise via ADMI_WINDOW_DAYS et sortie dans un dossier isolé
- Page /resultats: résumé, tableau paginé (15 lignes/page), téléchargement CSV & DOCX
- Gestion d'erreurs et messages clairs (flash)
- Couleurs FPM (#006b01) et police Tahoma
"""
import os
import io
import csv
import json
import shutil
import traceback
from datetime import datetime, date
from pathlib import Path
from subprocess import run, PIPE, CalledProcessError

from flask import (
    Flask, render_template, request, redirect, url_for, flash,
    send_from_directory, session
)

# --- Configuration de base ---
BASE_DIR = Path(__file__).resolve().parent
REPORTS_ROOT = BASE_DIR / "reports"  # chaque exécution aura un sous-dossier daté
DEFAULT_PROCESSOR = BASE_DIR / "fraude_detect_admi.py"  # script fourni par défaut

REPORTS_ROOT.mkdir(exist_ok=True, parents=True)

def create_app():
    app = Flask(__name__)
    app.secret_key = os.environ.get("FLASK_SECRET_KEY", "dev-secret-change-me")
    app.config["MAX_CONTENT_LENGTH"] = 16 * 1024 * 1024  # 16MB (pas utilisé ici, mais ok)
    return app

app = create_app()

# --------- Helpers ---------
def parse_date(s: str):
    try:
        return datetime.strptime(s, "%Y-%m-%d").date()
    except Exception:
        return None

def compute_days_window(d1: date, d2: date) -> int:
    # Inclusif: du début à la fin, minimum 1 jour
    return max(1, (d2 - d1).days + 1)

def build_run_folder() -> Path:
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    out = REPORTS_ROOT / f"run_{ts}"
    out.mkdir(parents=True, exist_ok=True)
    return out

def paginate(rows, page: int, per_page: int = 15):
    total = len(rows)
    start = (page - 1) * per_page
    end = start + per_page
    return rows[start:end], total

def generate_docx(summary: dict, csv_path: Path, out_docx: Path):
    # Génère un rapport Word simple avec python-docx
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Tahoma'
    style.font.size = Pt(10)

    doc.add_heading('Rapport de détection de fraudes - DICT / FPM', 0)
    doc.add_paragraph(f"Période analysée : {summary.get('debut')} au {summary.get('fin')}")
    doc.add_paragraph(f"Nombre total d'anomalies : {summary.get('nb_total', 0)}")

    doc.add_heading('Aperçu des anomalies', level=1)
    # Insère les 20 premières lignes en table
    import csv
    with csv_path.open("r", encoding="utf-8", newline="") as f:
        reader = csv.reader(f)
        try:
            headers = next(reader)
        except StopIteration:
            headers = []
            rows = []
        else:
            rows = [next(reader) for _ in range(20) if True]

    if headers:
        table = doc.add_table(rows=1, cols=len(headers))
        hdr_cells = table.rows[0].cells
        for j, h in enumerate(headers):
            hdr_cells[j].text = h

        for row in rows:
            row_cells = table.add_row().cells
            for j, val in enumerate(row):
                row_cells[j].text = str(val)

    doc.save(out_docx)

# --------- Routes ---------
@app.route("/")
def index():
    return render_template("index.html")

@app.route("/dates", methods=["GET", "POST"])
def dates():
    if request.method == "POST":
        d1 = parse_date(request.form.get("date_debut", ""))
        d2 = parse_date(request.form.get("date_fin", ""))

        if not d1 or not d2:
            flash("Veuillez saisir deux dates valides (AAAA-MM-JJ).", "error")
            return redirect(url_for("dates"))
        if d2 < d1:
            flash("La date de fin doit être supérieure ou égale à la date de début.", "error")
            return redirect(url_for("dates"))

        # Exécute le traitement en utilisant EXCLUSIVEMENT le script local
        try:
            out_folder = build_run_folder()
            # Calcule la fenêtre temporelle et mappe sur ADMI_WINDOW_DAYS
            days = compute_days_window(d1, d2)
            env = os.environ.copy()
            env["ADMI_WINDOW_DAYS"] = str(days)
            env["ADMI_OUT_DIR"] = str(out_folder)  # contiendra auto_all.csv

            proc = run(
                [os.environ.get("PYTHON_BIN", "python"), str(DEFAULT_PROCESSOR)],
                cwd=str(BASE_DIR),
                stdout=PIPE, stderr=PIPE, text=True, env=env, timeout=None
            )
            log_txt = (out_folder / "execution.log")
            log_txt.write_text(proc.stdout + "\n\n[STDERR]\n" + proc.stderr, encoding="utf-8")

            if proc.returncode != 0:
                raise CalledProcessError(proc.returncode, proc.args, proc.stdout, proc.stderr)

            csv_all = out_folder / "auto_all.csv"
            if not csv_all.exists():
                # Cherche un CSV alternatif si le principal n'est pas présent
                candidates = list(out_folder.glob("auto_*.csv"))
                if not candidates:
                    raise RuntimeError("Aucun fichier de résultats n'a été généré.")
                csv_all = candidates[0]

            # Lis le CSV pour le résumé
            import pandas as pd
            df = pd.read_csv(csv_all)
            nb_total = len(df.index)

            # Résumé en session
            session["summary"] = {
                "debut": d1.isoformat(),
                "fin": d2.isoformat(),
                "nb_total": int(nb_total),
                "folder": str(out_folder),
                "csv": str(csv_all.name)
            }

            # Génère DOCX
            try:
                docx_path = out_folder / "rapport_fraudes.docx"
                generate_docx(session["summary"], csv_all, docx_path)
            except Exception as e:
                (out_folder / "docx_error.txt").write_text(str(e), encoding="utf-8")

            return redirect(url_for("resultats"))

        except Exception as e:
            traceback_txt = traceback.format_exc()
            flash(f"Erreur pendant le traitement: {e}", "error")
            # Log erreur globale
            (out_folder / "fatal_error.txt").write_text(traceback_txt, encoding="utf-8")
            return redirect(url_for("dates"))

    # GET
    return render_template("dates.html")

@app.route("/resultats")
def resultats():
    summary = session.get("summary")
    if not summary:
        flash("Aucun résultat disponible. Veuillez relancer un traitement.", "error")
        return redirect(url_for("index"))

    folder = Path(summary["folder"])
    csv_name = summary.get("csv", "auto_all.csv")
    csv_path = folder / csv_name

    # Lecture complète pour pagination
    headers = []
    rows = []
    if csv_path.exists():
        with csv_path.open("r", encoding="utf-8", newline="") as f:
            reader = csv.reader(f)
            try:
                headers = next(reader)
            except StopIteration:
                headers = []
            else:
                rows = list(reader)

    page = max(1, int(request.args.get("page", "1")))
    per_page = 15
    page_rows, total = paginate(rows, page, per_page)
    last_page = max(1, (total + per_page - 1) // per_page)

    return render_template(
        "resultats.html",
        summary=summary,
        headers=headers,
        rows=page_rows,
        page=page,
        last_page=last_page,
        csv_file=csv_name,
        docx_file="rapport_fraudes.docx",
        run_folder=folder.name
    )

@app.route("/download/<run_id>/<path:filename>")
def download(run_id, filename):
    # Protection basique contre path traversal
    safe_name = os.path.basename(filename)
    run_folder = REPORTS_ROOT / run_id
    return send_from_directory(run_folder, safe_name, as_attachment=True)

@app.errorhandler(413)
def too_large(e):
    flash("Fichier trop volumineux.", "error")
    return redirect(url_for("dates"))

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000, debug=True)
